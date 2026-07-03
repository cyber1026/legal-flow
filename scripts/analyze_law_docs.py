#!/usr/bin/env python3
"""Analyze structure consistency of law DOCX files."""
from __future__ import annotations

import json
import re
import zipfile
from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

REPO_ROOT = Path(__file__).resolve().parents[1]
LAW_DIR = REPO_ROOT / "law"

# Chinese legal structure patterns
PATTERNS = {
    "title_like": re.compile(r"^中华人民共和国.+法$"),
    "part": re.compile(r"^第[一二三四五六七八九十百千零〇]+编\s"),
    "chapter": re.compile(r"^第[一二三四五六七八九十百千零〇]+章\s"),
    "section": re.compile(r"^第[一二三四五六七八九十百千零〇]+节\s"),
    "article": re.compile(r"^第[一二三四五六七八九十百千零〇百零]+条\s"),
    "article_alt": re.compile(r"^第\d+条\s"),
    "book_volume": re.compile(r"^第[一二三四五六七八九十百千零〇]+分编\s"),
    "preface": re.compile(r"^(总则|附则|序言|前言)$"),
    "effective_date": re.compile(r"(自\s*\d{4}年\d{1,2}月\d{1,2}日|本法自.+起施行)"),
    "revision_note": re.compile(r"(根据|已由|自.+修正|修订)"),
}


def get_paragraph_text(p) -> str:
    return (p.text or "").strip()


def get_paragraph_style(p) -> str:
    try:
        return p.style.name if p.style else "None"
    except Exception:
        return "Unknown"


def analyze_docx(path: Path) -> dict:
    doc = Document(path)
    paragraphs = [get_paragraph_text(p) for p in doc.paragraphs if get_paragraph_text(p)]
    styles = [get_paragraph_style(p) for p in doc.paragraphs if get_paragraph_text(p)]

    pattern_counts = {k: 0 for k in PATTERNS if k not in ("title_like",)}
    matched_lines = defaultdict(list)

    for i, text in enumerate(paragraphs):
        for name, pat in PATTERNS.items():
            if pat.search(text):
                if name not in ("title_like",):
                    pattern_counts[name] = pattern_counts.get(name, 0) + 1
                if name in ("chapter", "section", "article", "part", "book_volume"):
                    matched_lines[name].append((i, text[:80]))

    # Tables
    table_count = len(doc.tables)
    table_shapes = []
    for t in doc.tables:
        rows = len(t.rows)
        cols = len(t.columns) if t.rows else 0
        table_shapes.append((rows, cols))

    # Core properties
    cp = doc.core_properties
    props = {
        "title": cp.title or "",
        "author": cp.author or "",
        "subject": cp.subject or "",
        "created": str(cp.created) if cp.created else "",
        "modified": str(cp.modified) if cp.modified else "",
    }

    # Filename metadata
    stem = path.stem
    date_match = re.search(r"_(\d{8})$", stem)
    law_name_from_file = re.sub(r"_\d{8}$", "", stem)

    # First N non-empty paragraphs (structure header)
    header_preview = paragraphs[:15]

    # Style distribution
    style_counter = Counter(styles)

    # Detect if title is first paragraph
    first_para = paragraphs[0] if paragraphs else ""
    has_title_first = bool(PATTERNS["title_like"].match(first_para)) or law_name_from_file in first_para

    # Article numbering style
    article_samples = [t for t in paragraphs if PATTERNS["article"].match(t)][:3]
    article_alt_count = sum(1 for t in paragraphs if PATTERNS["article_alt"].match(t))

    return {
        "file": path.name,
        "law_name_file": law_name_from_file,
        "version_date": date_match.group(1) if date_match else None,
        "paragraph_count": len(paragraphs),
        "non_empty_para_in_doc": sum(1 for p in doc.paragraphs if get_paragraph_text(p)),
        "table_count": table_count,
        "table_shapes": table_shapes,
        "core_properties": props,
        "style_distribution": dict(style_counter.most_common(10)),
        "unique_styles": sorted(set(styles)),
        "pattern_counts": pattern_counts,
        "has_title_first": has_title_first,
        "first_para": first_para[:100],
        "article_alt_count": article_alt_count,
        "article_samples": [s[:60] for s in article_samples],
        "chapter_samples": [t[:60] for _, t in matched_lines.get("chapter", [])[:3]],
        "section_samples": [t[:60] for _, t in matched_lines.get("section", [])[:3]],
        "part_samples": [t[:60] for _, t in matched_lines.get("part", [])[:3]],
        "header_preview": [h[:80] for h in header_preview],
        "avg_para_len": sum(len(p) for p in paragraphs) / len(paragraphs) if paragraphs else 0,
        "max_para_len": max((len(p) for p in paragraphs), default=0),
    }


def docx_xml_signature(path: Path) -> dict:
    """Inspect raw OOXML structure inside docx zip."""
    sig = {"has_styles_xml": False, "has_numbering_xml": False, "body_child_tags": Counter()}
    try:
        with zipfile.ZipFile(path) as z:
            names = z.namelist()
            sig["has_styles_xml"] = "word/styles.xml" in names
            sig["has_numbering_xml"] = "word/numbering.xml" in names
            if "word/document.xml" in names:
                xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
                for tag in ("w:p", "w:tbl", "w:sectPr", "w:br", "w:hyperlink"):
                    sig["body_child_tags"][tag] = xml.count(f"<{tag}")
    except Exception as e:
        sig["error"] = str(e)
    return sig


def main():
    files = sorted(LAW_DIR.glob("*.docx"))
    results = []
    xml_sigs = []

    for f in files:
        info = analyze_docx(f)
        info["xml"] = docx_xml_signature(f)
        results.append(info)
        xml_sigs.append(info["xml"])

    # Aggregate statistics
    n = len(results)

    # Format
    formats = Counter(f.suffix for f in files)

    # Pattern stats across docs
    pattern_matrix = defaultdict(list)
    for r in results:
        for k, v in r["pattern_counts"].items():
            pattern_matrix[k].append(v)

    # Structure fingerprint per doc
    def fingerprint(r):
        return (
            tuple(sorted(r["unique_styles"])),
            r["table_count"] > 0,
            r["pattern_counts"].get("part", 0) > 0,
            r["pattern_counts"].get("book_volume", 0) > 0,
            r["pattern_counts"].get("chapter", 0) > 0,
            r["pattern_counts"].get("section", 0) > 0,
            r["pattern_counts"].get("article", 0) > 0,
        )

    fps = [fingerprint(r) for r in results]
    fp_counter = Counter(fps)

    # Style sets
    all_styles = set()
    for r in results:
        all_styles.update(r["unique_styles"])
    style_presence = {s: sum(1 for r in results if s in r["unique_styles"]) for s in sorted(all_styles)}

    # Core properties consistency
    prop_keys = ["title", "author", "subject"]
    prop_values = {k: Counter(r["core_properties"][k] for r in results) for k in prop_keys}

    # Paragraph count range
    para_counts = [r["paragraph_count"] for r in results]

    # Article counts
    article_counts = [r["pattern_counts"].get("article", 0) for r in results]

    # Header structure comparison (first 5 lines pattern types)
    def classify_line(text):
        for name, pat in PATTERNS.items():
            if pat.search(text):
                return name
        if not text:
            return "empty"
        if len(text) < 30 and "法" in text:
            return "short_title"
        return "body"

    header_patterns = []
    for r in results:
        types = [classify_line(h) for h in r["header_preview"][:8]]
        header_patterns.append(tuple(types))

    header_fp = Counter(header_patterns)

    report = {
        "summary": {
            "total_files": n,
            "file_formats": dict(formats),
            "all_same_format": len(formats) == 1 and ".docx" in formats,
            "paragraph_count_min": min(para_counts),
            "paragraph_count_max": max(para_counts),
            "paragraph_count_avg": round(sum(para_counts) / n, 1),
            "article_count_min": min(article_counts),
            "article_count_max": max(article_counts),
            "structure_fingerprint_groups": len(fp_counter),
            "structure_all_identical": len(fp_counter) == 1,
            "unique_style_names_total": len(all_styles),
            "files_with_tables": sum(1 for r in results if r["table_count"] > 0),
            "files_with_part": sum(1 for r in results if r["pattern_counts"].get("part", 0) > 0),
            "files_with_section": sum(1 for r in results if r["pattern_counts"].get("section", 0) > 0),
            "files_with_chapter": sum(1 for r in results if r["pattern_counts"].get("chapter", 0) > 0),
            "files_with_article": sum(1 for r in results if r["pattern_counts"].get("article", 0) > 0),
            "files_with_arabic_articles": sum(1 for r in results if r["article_alt_count"] > 0),
            "has_styles_xml_all": all(x.get("has_styles_xml") for x in xml_sigs),
            "has_numbering_xml_all": all(x.get("has_numbering_xml") for x in xml_sigs),
        },
        "pattern_stats_per_doc": {
            k: {"min": min(v), "max": max(v), "docs_with_zero": sum(1 for x in v if x == 0)}
            for k, v in pattern_matrix.items()
        },
        "structure_fingerprints": [
            {"count": c, "fingerprint": str(fp)} for fp, c in fp_counter.most_common()
        ],
        "style_presence": style_presence,
        "core_properties": {k: dict(v) for k, v in prop_values.items()},
        "header_pattern_groups": len(header_fp),
        "per_file": results,
    }

    out_path = REPO_ROOT / "scripts" / "law_analysis_report.json"
    out_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")

    # Print human-readable summary
    print("=" * 60)
    print("法律文档格式与结构一致性分析报告")
    print("=" * 60)
    s = report["summary"]
    print(f"\n【文件格式】共 {s['total_files']} 个文件")
    print(f"  格式: {s['file_formats']}")
    print(f"  格式是否统一: {'是' if s['all_same_format'] else '否'}")

    print(f"\n【OOXML 底层结构】")
    print(f"  均含 styles.xml: {s['has_styles_xml_all']}")
    print(f"  均含 numbering.xml: {s['has_numbering_xml_all']}")

    print(f"\n【段落规模】")
    print(f"  非空段落数: {s['paragraph_count_min']} ~ {s['paragraph_count_max']} (均值 {s['paragraph_count_avg']})")
    print(f"  法条(第X条)匹配数: {s['article_count_min']} ~ {s['article_count_max']}")

    print(f"\n【法律层级结构出现率】")
    print(f"  含「编」: {s['files_with_part']}/{n}")
    print(f"  含「章」: {s['files_with_chapter']}/{n}")
    print(f"  含「节」: {s['files_with_section']}/{n}")
    print(f"  含「条」(中文数字): {s['files_with_article']}/{n}")
    print(f"  含阿拉伯数字条: {s['files_with_arabic_articles']}/{n}")
    print(f"  含表格: {s['files_with_tables']}/{n}")

    print(f"\n【结构指纹】共 {s['structure_fingerprint_groups']} 种结构类型")
    print(f"  结构完全一致: {'是' if s['structure_all_identical'] else '否'}")

    print(f"\n【Word 样式】共 {s['unique_style_names_total']} 种不同样式名")
    for st, cnt in sorted(style_presence.items(), key=lambda x: -x[1])[:15]:
        print(f"  {st!r}: {cnt}/{n} 文件")

    print(f"\n【文档元数据 core_properties】")
    for k, vals in report["core_properties"].items():
        print(f"  {k}: {vals}")

    print(f"\n【各文件明细】")
    print(f"{'文件名':<45} {'段落':>6} {'章':>4} {'节':>4} {'条':>4} {'表':>3} {'样式数':>4}")
    print("-" * 80)
    for r in results:
        pc = r["pattern_counts"]
        print(
            f"{r['file']:<45} {r['paragraph_count']:>6} "
            f"{pc.get('chapter',0):>4} {pc.get('section',0):>4} {pc.get('article',0):>4} "
            f"{r['table_count']:>3} {len(r['unique_styles']):>4}"
        )

    print(f"\n【首部结构示例对比】(前3个文件)")
    for r in results[:3]:
        print(f"\n--- {r['law_name_file']} ---")
        for i, line in enumerate(r["header_preview"][:8]):
            print(f"  [{i}] {line}")

    print(f"\n完整 JSON 报告: {out_path}")


if __name__ == "__main__":
    main()
